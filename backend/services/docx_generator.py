from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re


class DocxGenerator:
    """Service for generating ATS-friendly DOCX files"""

    @staticmethod
    def generate_docx(optimized_resume_text, candidate_name="Resume"):
        """
        Generate a DOCX file from optimized resume text

        Args:
            optimized_resume_text (str): Optimized resume text
            candidate_name (str): Candidate name for filename

        Returns:
            BytesIO: Binary DOCX file data

        Raises:
            Exception: If generation fails
        """
        try:
            # Create new document
            doc = Document()

            # Apply ATS-friendly formatting
            DocxGenerator._apply_document_formatting(doc)

            # Parse and add content
            DocxGenerator._add_resume_content(doc, optimized_resume_text)

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return file_stream

        except Exception as e:
            raise Exception(f"DOCX generation failed: {str(e)}")

    @staticmethod
    def _apply_document_formatting(doc):
        """Apply ATS-friendly document formatting"""

        # Set page margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)

        # Set default font for Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)

        # Configure Heading 1 style (for name)
        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 0)

        # Configure Heading 2 style (for section headers)
        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(12)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _add_resume_content(doc, text):
        """Parse resume text and add to document with proper formatting"""

        lines = text.split('\n')
        current_section = None
        is_first_line = True

        for line in lines:
            line = line.strip()

            if not line:
                # Empty line - add spacing
                doc.add_paragraph()
                continue

            # Check if this is a section header (all caps or specific patterns)
            if DocxGenerator._is_section_header(line):
                # Add section header
                p = doc.add_heading(line, level=2)
                current_section = line.lower()
                continue

            # Check if this is the candidate name (first non-empty line)
            if is_first_line:
                p = doc.add_heading(line, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                is_first_line = False
                continue

            # Check if this is a bullet point
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Add bullet point
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)

            # Apply formatting to paragraph
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    @staticmethod
    def _is_section_header(line):
        """Determine if a line is a section header"""

        # Check if line is all caps (likely a section header)
        if line.isupper() and len(line) > 2:
            return True

        # Check for common section headers
        section_keywords = [
            'CONTACT INFORMATION', 'PROFESSIONAL SUMMARY', 'SUMMARY',
            'CORE COMPETENCIES', 'SKILLS', 'TECHNICAL SKILLS',
            'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EXPERIENCE',
            'EDUCATION', 'CERTIFICATIONS', 'CERTIFICATES',
            'PROJECTS', 'KEY PROJECTS', 'AWARDS', 'ACHIEVEMENTS'
        ]

        line_upper = line.upper()
        for keyword in section_keywords:
            if keyword in line_upper:
                return True

        return False
